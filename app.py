import streamlit as st
from PIL import Image
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from textblob import TextBlob
import helper
import preprocessor
import openai
from auth_pages import show_sign_up_page, show_login_page
from io import BytesIO  # Import BytesIO
from docx import Document
from io import StringIO
from docx.shared import Inches
from datetime import datetime



def get_chatbot_response(user_input):
    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=user_input,
        max_tokens=150
    )
    return response.choices[0].text.strip()


def main():
    st.title("WhatsApp Chat Analyzer with Chatbot")

    # Chatbot Section
    st.header("Chatbot Assistant")
    st.write("Ask me anything about the app!")

    # Text input for user query
    user_query = st.text_input("Your question:")

    if st.button("Ask Chatbot"):
        if user_query:
            response = get_chatbot_response(user_query)
            st.write(f"Chatbot: {response}")
        else:
            st.warning("Please enter a question.")



# Function to save Plotly figures as images
def save_plot_as_image(fig, filename):
    fig.write_image(filename, format='png')


# Function to generate Word report
def generate_word_report(num_messages, words, num_media_messages, num_links, most_used_words_df, emoji_df,
                         sentiment_counts,
                         timeline_img, daily_timeline_img, most_used_words_img, emoji_img, sentiment_img,
                         most_active_days_img):
    doc = Document()
    doc.add_heading('Chat Analysis Report', 0)

    # Add basic stats to Word
    doc.add_heading('Top Chat Statistics', level=1)
    doc.add_paragraph(f'Total Messages: {num_messages}')
    doc.add_paragraph(f'Total Words: {words}')
    doc.add_paragraph(f'Media Shared: {num_media_messages}')
    doc.add_paragraph(f'Links Shared: {num_links}')

    # Add charts to Word
    doc.add_heading('Monthly Timeline', level=2)
    doc.add_picture(timeline_img, width=Inches(6))

    doc.add_heading('Daily Timeline', level=2)
    doc.add_picture(daily_timeline_img, width=Inches(6))

    doc.add_heading('Top 10 Most Used Words', level=2)
    doc.add_picture(most_used_words_img, width=Inches(6))

    doc.add_heading('Top Emojis Used', level=2)
    doc.add_picture(emoji_img, width=Inches(6))

    doc.add_heading('Sentiment Analysis', level=2)
    doc.add_picture(sentiment_img, width=Inches(6))

    doc.add_heading('Most Active Days', level=2)
    doc.add_picture(most_active_days_img, width=Inches(6))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Function to generate CSV data
def generate_csv_report(num_messages, words, num_media_messages, num_links, most_used_words_df, emoji_df,
                        sentiment_counts):
    output = StringIO()

    # Write basic stats into the CSV
    output.write('Top Chat Statistics\n')
    output.write(f'Total Messages, {num_messages}\n')
    output.write(f'Total Words, {words}\n')
    output.write(f'Media Shared, {num_media_messages}\n')
    output.write(f'Links Shared, {num_links}\n')

    # Most Used Words section
    output.write('\nTop 10 Most Used Words\n')
    most_used_words_df.head(10).to_csv(output, index=False)

    # Emoji Analysis section
    output.write('\nTop Emojis Used\n')
    emoji_df.head().to_csv(output, index=False)

    # Sentiment Analysis section
    output.write('\nSentiment Analysis\n')
    sentiment_counts.to_csv(output, index=True, header=False)

    return output.getvalue()





def analyze_sentiments(df):
    """Function to analyze sentiment of the messages"""

    def get_sentiment(text):
        analysis = TextBlob(text)
        if analysis.sentiment.polarity > 0:
            return 'Positive'
        elif analysis.sentiment.polarity < 0:
            return 'Negative'
        else:
            return 'Neutral'

    df['sentiment'] = df['message'].apply(get_sentiment)
    sentiment_counts = df['sentiment'].value_counts()
    return sentiment_counts


def get_most_used_words(df):
    """Function to get the most used words from the messages"""
    from collections import Counter
    import re

    # Clean and tokenize the text
    text = ' '.join(df['message'].tolist())
    words = re.findall(r'\b\w+\b', text.lower())
    word_counts = Counter(words)

    # Convert to DataFrame
    most_used_words_df = pd.DataFrame(word_counts.items(), columns=['Word', 'Count'])
    most_used_words_df = most_used_words_df.sort_values(by='Count', ascending=False)

    return most_used_words_df


def get_most_active_days(df):
    """Function to get the most active days"""
    # Group by date and count messages
    most_active_days_df = df.groupby('date').size().reset_index(name='message_count')
    most_active_days_df = most_active_days_df.sort_values(by='message_count', ascending=False)

    return most_active_days_df


def load_css(css_path):
    with open(css_path, "r") as f:
        css = f.read()
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)



# Function to load custom CSS
def load_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Simple rule-based chatbot function
def rule_based_chatbot(user_input):
    # Define some simple rules
    user_input = user_input.lower()  # Convert user input to lowercase for easier matching

    if "hello" in user_input or "hi" in user_input:
        return "Hello! How can I assist you with the WhatsApp Chat Analyzer today?"
    elif "analyze" in user_input or "analysis" in user_input:
        return "You can upload your WhatsApp chat file to start the analysis on the 'Analysis' page."
    elif "signup" in user_input or "sign up" in user_input:
        return "You can create a new account on the Sign-up page."
    elif "login" in user_input:
        return "You can log in on the Login page."
    elif "feedback" in user_input:
        return "We appreciate your feedback! Please use the sidebar to submit it."
    elif "upload" in user_input:
        return "To upload your chat, go to the 'Analysis' page and use the file uploader."
    elif "insights" in user_input or "insight" in user_input:
        return "The WhatsApp Chat Analyzer provides insights on sentiment, user activity, and engagement metrics."
    elif "sentiment" in user_input:
        return "You can analyze the sentiment of your chats to understand the emotional tone—positive, negative, or neutral."
    elif "user activity" in user_input:
        return "You can view user activity metrics such as message count, response time, and engagement levels."
    elif "most used words" in user_input:
        return "The tool allows you to see the most frequently used words in your chats, providing insights into key topics."
    elif "emoji analysis" in user_input:
        return "You can analyze emoji usage to see which emojis are most commonly used in your chats."
    elif "export" in user_input:
        return "To analyze your chats, first, export your WhatsApp chat without media and then upload it here."
    elif "help" in user_input or "assistance" in user_input:
        return "I’m here to help! You can ask me about chat analysis, user metrics, or how to navigate the application."
    elif "reports" in user_input or "generate report" in user_input:
        return "You can generate detailed reports on your chat analysis after uploading the file."
    elif "privacy" in user_input or "data security" in user_input:
        return "We prioritize your privacy. Your chat data is only used for analysis and not stored."
    elif "troubleshoot" in user_input or "issue" in user_input:
        return "If you encounter issues, please ensure your file is in the correct format and without media."
    elif "features" in user_input:
        return "Our features include sentiment analysis, user activity tracking, most used words, and emoji analysis."
    elif "user metrics" in user_input:
        return "You can view metrics for each user in your chats, helping you understand interactions better."
    elif "data visualization" in user_input:
        return "The application provides visualizations like charts and graphs for easy understanding of your chat data."
    elif "custom analysis" in user_input:
        return "We are working on providing customizable analysis options to suit your specific needs."
    elif "navigate" in user_input:
        return "You can navigate through the application using the sidebar to access different pages."
    elif "contact support" in user_input or "need help" in user_input:
        return "If you need further assistance, please reach out through our contact page."
    else:
        return "I'm sorry, I didn't understand that. Can you please rephrase?"
# Main function
def main():
    load_css("style.css")

    # Sidebar navigation using buttons
    if st.sidebar.button("Home"):
        st.query_params.update({"page": "home"})
    if st.sidebar.button("Analysis"):
        st.query_params.update({"page": "analysis"})
    if st.sidebar.button("Sign up"):
        st.query_params.update({"page": "signup"})
    if st.sidebar.button("Login"):
        st.query_params.update({"page": "login"})

    # Get the page from the query params
    query_params = st.query_params
    page = query_params.get("page", "home")  # Default to 'home' if no param is found

    if page != "analysis":
        # Add image at the bottom of the sidebar for pages other than 'analysis'
        st.sidebar.markdown("---")  # Line separator
        st.sidebar.image(Image.open('wca.jpeg'), width=300)

    # Main content based on the selected page
    if page == "home":
        st.image("images/logo.png", use_column_width=True)
        st.title('WhatsApp Chat Analyzer')
        st.write("""Welcome to **WhatsApp Chat Analyzer** - Your Personal WhatsApp Chat Insight Engine!
        Dive deep into your chat narratives and discover patterns you never noticed before. 
        From sentiment landscapes to user engagement metrics, unveils a spectrum of data-driven revelations from your chats.""")
        st.write('**Guidelines to Begin:**')
        st.write("""
           1. Export your WhatsApp chat Without Media.
           2. Upload the exported chat file.
           3. Choose the user and the type of analysis.
           4. View insights and charts based on your chat data.
           """)

        # Rule-based Chatbot section on the home page
        st.header("Chatbot Assistant")
        st.write("Ask me anything about the app!")
        user_query = st.text_input("Your question:")
        if st.button("Ask Chatbot"):
            if user_query:
                response = rule_based_chatbot(user_query)
                st.write(f"Chatbot: {response}")
            else:
                st.warning("Please enter a question.")

        # Footer section
        st.markdown("""<footer style="background-color: #333; color: white; text-align: center; padding: 1rem; width: 100%; position: relative;">
            <p>&copy; 2024 WhatsApp Chat Analyzer. All rights reserved.</p>
            </footer>""", unsafe_allow_html=True)

        st.sidebar.markdown("---")
        st.sidebar.header("We Value Your Feedback")

        # Dropdown for selecting feedback
        was_helpful = st.sidebar.selectbox("Did you find our insights useful?",
                                           ["Please choose an option", "Yes, very helpful", "Somewhat helpful", "Not helpful"])
        if was_helpful != "Please choose an option":
            feedback = st.sidebar.text_area("Kindly share any additional comments or suggestions...")
            if st.sidebar.button("Submit Feedback"):
                st.sidebar.success("Thank you for sharing your thoughts!")

    elif page == "signup":
        show_sign_up_page()
    elif page == "login":
        show_login_page()


    elif page == "analysis":
        # File uploader in the sidebar
        st.sidebar.markdown("---")  # Line separator
        uploaded_file = st.sidebar.file_uploader("Choose a file")

        # Check if a file has been uploaded
        if uploaded_file:
            bytes_data = uploaded_file.getvalue()
            data = bytes_data.decode("utf-8")
            df = preprocessor.preprocess(data)  # Assuming preprocessor has a method named preprocess

            # Display a message indicating analysis generation
            st.markdown('<p style="font-family:Roboto; color:#ffa81a; font-size: 20px; font-weight: bold">Here is your analysis generated:</p>', unsafe_allow_html=True)

            # Fetch unique users and create a user dropdown
            user_list = df['user'].unique().tolist()
            if 'group_notification' in user_list:
                user_list.remove('group_notification')
            user_list.sort()
            user_list.insert(0, 'Overall')  # Add 'Overall' option for all users

            # Sidebar selectbox for user selection
            selected_user = st.sidebar.selectbox("Show analysis with respect to", user_list)

            # Further analysis logic based on the selected user can be added here


            # Visualization selection
            visualization_option = st.sidebar.selectbox(
                "Select Visualization",
                ["Select", "Chat Data", "Top Chat Statistics", "Monthly Timeline", "Daily Timeline",
                 "Most Used Words", "Most Active Users", "WordCloud", "Emoji Analysis", "Sentiment Analysis", "Most Active Days", "All Analysis"]
            )

            if st.sidebar.button("Show Analysis"):
                if visualization_option == "Chat Data":
                    # Show raw chat data
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Chat Data</p></u>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">Here is the raw chat data you uploaded:</p>',
                        unsafe_allow_html=True
                    )
                    st.dataframe(df)

                elif visualization_option == "Top Chat Statistics":
                    # Stats Area
                    num_messages, words, num_media_messages, num_links = helper.fetch_stats(selected_user, df)

                    st.markdown(
                        '<u><p style="font-family:Roboto; color:#5e76c4; font-size: 50px; font-weight: bold">Top Chat Statistics</p></u>',
                        unsafe_allow_html=True
                    )
                    col1, col2, col3, col4 = st.columns(4)

                    with col1:
                        st.markdown(
                            '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Total Messages</p>',
                            unsafe_allow_html=True
                        )
                        st.title(num_messages)
                    with col2:
                        st.markdown(
                            '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Total Words</p>',
                            unsafe_allow_html=True
                        )
                        st.title(words)
                    with col3:
                        st.markdown(
                            '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Media Shared</p>',
                            unsafe_allow_html=True
                        )
                        st.title(num_media_messages)
                    with col4:
                        st.markdown(
                            '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Links Shared</p>',
                            unsafe_allow_html=True
                        )
                        st.title(num_links)

                elif visualization_option == "Monthly Timeline":

                    # Sample DataFrame for monthly timeline analysis (replace with your helper function)
                    # Assuming the 'time' column is in 'YYYY-MM-DD' format
                    def mock_monthly_timeline(selected_user, df):
                        return pd.DataFrame({
                            'time': pd.date_range(start='2022-01-01', periods=24, freq='M'),  # Mock dates
                            'message': [i * 10 for i in range(24)]  # Mock message counts
                        })

                    # Replace with your actual data and function
                    df = pd.DataFrame()  # Placeholder
                    selected_user = "User"  # Placeholder
                    timeline = mock_monthly_timeline(selected_user, df)

                    # Streamlit page structure
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Monthly Timeline</p></u>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This status shows total number of chats per month in group or personal.</p>',
                        unsafe_allow_html=True
                    )

                    # Date input for selecting a full start and end date
                    start_date = st.date_input("Select start date (default day is 1)", value=datetime(2023, 1, 1))
                    end_date = st.date_input("Select end date (default day is 1)", value=datetime(2023, 12, 1))

                    # Extract month and year from the selected dates
                    start_month_year = datetime(start_date.year, start_date.month, 1)
                    end_month_year = datetime(end_date.year, end_date.month, 1)

                    # Convert 'time' column to datetime if not already in that format
                    timeline['time'] = pd.to_datetime(timeline['time'])

                    # Filter based on the selected start and end month-year
                    filtered_timeline = timeline[
                        (timeline['time'] >= start_month_year) & (timeline['time'] <= end_month_year)]

                    # Plot the filtered data using bar chart
                    fig = px.bar(
                        filtered_timeline,
                        x='time',
                        y='message',
                        title='Monthly Chats',
                        labels={'time': 'Months', 'message': 'Number of Chats'},
                        template='plotly_dark'
                    )
                    st.plotly_chart(fig)

                elif visualization_option == "Daily Timeline":

                    # Assuming helper.daily_timeline returns a DataFrame with 'only_date' and 'message' columns
                    daily_timeline = helper.daily_timeline(selected_user, df)

                    # Convert 'only_date' to datetime with flexible parsing
                    # Use 'dayfirst=True' to prioritize day-first format, and 'errors' to coerce invalid dates
                    daily_timeline['only_date'] = pd.to_datetime(daily_timeline['only_date'], format='mixed',
                                                                 dayfirst=True, errors='coerce')

                    # Drop rows with invalid dates (if any)
                    daily_timeline = daily_timeline.dropna(subset=['only_date'])

                    # Ensure the timeline is sorted by date for proper plotting
                    daily_timeline = daily_timeline.sort_values(by='only_date')

                    # Add date range filter
                    st.markdown(
                        '<br><br><u><p style="font-family:georgia; color:#ca124d; font-size: 40px; font-weight: bold">Daily Timeline</p></u>',
                        unsafe_allow_html=True)
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This status shows the total number of chats per day in group or personal.</p>',
                        unsafe_allow_html=True)

                    # Date input widgets for selecting range
                    start_date = st.date_input("Select start date", value=daily_timeline['only_date'].min().date())
                    end_date = st.date_input("Select end date", value=daily_timeline['only_date'].max().date())

                    # Filter the DataFrame based on the selected date range
                    filtered_timeline = daily_timeline[(daily_timeline['only_date'] >= pd.to_datetime(start_date)) &
                                                       (daily_timeline['only_date'] <= pd.to_datetime(end_date))]

                    # Plotting the filtered data as a bar graph
                    fig = px.bar(filtered_timeline, x='only_date', y='message', title='Daily Chats',
                                 labels={'only_date': 'Date', 'message': 'Number of Chats'}, template='plotly_dark')

                    # Display the plot
                    st.plotly_chart(fig)

                elif visualization_option == "Most Used Words":
                    # Most Used Words using Plotly
                    most_used_words_df = get_most_used_words(df)
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Used Words</p></u>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">Here are the most frequently used words in your chat.</p>',
                        unsafe_allow_html=True
                    )
                    fig = px.bar(most_used_words_df.head(10), x='Word', y='Count', title='Top 10 Most Used Words',
                                 labels={'Word': 'Words', 'Count': 'Frequency'}, template='plotly_dark')
                    st.plotly_chart(fig)

                elif visualization_option == "Most Active Users":
                    # Convert date inputs to datetime64[ns]
                    start_date = pd.to_datetime(st.date_input('Start Date', pd.to_datetime(df['date'].min())))
                    end_date = pd.to_datetime(st.date_input('End Date', pd.to_datetime(df['date'].max())))

                    # Ensure 'date' column in DataFrame is of type datetime64[ns]
                    df['date'] = pd.to_datetime(df['date'])

                    # Filter the DataFrame based on the selected date range
                    filtered_df = df[(df['date'] >= start_date) & (df['date'] <= end_date)]

                    if selected_user == 'Overall':
                        st.markdown(
                            '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Active Users</p></u>',
                            unsafe_allow_html=True
                        )
                        st.markdown(
                            '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This stats shows the most active users in the group.</p>',
                            unsafe_allow_html=True
                        )

                        # Pass the filtered DataFrame to your helper function
                        x, new_df = helper.most_active_users(filtered_df)

                        fig = px.bar(x, x.index, x.values, title='Top 5 Active Users',
                                     labels={'x': 'Users', 'y': 'Number of Chats'}, template='plotly_dark')
                        st.plotly_chart(fig)

                        st.dataframe(new_df)


                elif visualization_option == "WordCloud":
                    # WordCloud
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">WordCloud</p></u>',
                        unsafe_allow_html=True
                    )
                    df_wc = helper.create_wordcloud(selected_user, df)
                    fig = px.imshow(df_wc, title='WordCloud', labels={'x': '', 'y': ''}, template='plotly_dark')
                    st.plotly_chart(fig)

                elif visualization_option == "Emoji Analysis":
                    # Emoji analysis using Plotly
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Emoji Analysis</p></u>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This stats show most used emoji in the chats.</p>',
                        unsafe_allow_html=True
                    )
                    emoji_df = helper.emojis(selected_user, df)
                    col1, col2 = st.columns(2)

                    with col1:
                        st.dataframe(emoji_df)

                    with col2:
                        fig = go.Figure(data=[go.Pie(
                            labels=emoji_df['Emoji'].head(),  # Adjust column name
                            values=emoji_df['Count'].head(),  # Adjust column name
                            hole=0.3
                        )])
                        fig.update_layout(title="Emoji Distribution")
                        st.plotly_chart(fig)

                elif visualization_option == "Sentiment Analysis":
                    # Sentiment Analysis using Plotly
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Sentiment Analysis</p></u>',
                        unsafe_allow_html=True
                    )
                    sentiment_counts = analyze_sentiments(df)
                    fig = px.pie(sentiment_counts, values='count', names=sentiment_counts.index,
                                 title='Sentiment Analysis Distribution', template='plotly_dark')
                    st.plotly_chart(fig)

                elif visualization_option == "Most Active Days":
                    # Most Active Days Bar Chart using Plotly
                    st.markdown(
                        '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Active Days</p></u>',
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This bar chart displays the days with the highest message count.</p>',
                        unsafe_allow_html=True
                    )
                    most_active_days_df = get_most_active_days(df)
                    fig = px.bar(most_active_days_df.head(10), x='date', y='message_count',
                                 title='Top 10 Most Active Days',
                                 labels={'date': 'Date', 'message_count': 'Message Count'}, template='plotly_dark')
                    st.plotly_chart(fig)



                 #all anaylsis page
                elif visualization_option == "All Analysis":
                    if visualization_option == "All Analysis":
                        # Stats Area
                        num_messages, words, num_media_messages, num_links = helper.fetch_stats(selected_user, df)

                        st.markdown(
                            '<u><p style="font-family:Roboto; color:#5e76c4; font-size: 50px; font-weight: bold">Top Chat Statistics</p></u>',
                            unsafe_allow_html=True
                        )

                        col1, col2, col3, col4 = st.columns([1, 1, 1, 1])  # Equal-width columns

                        with col1:
                            st.markdown(
                                '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Total Messages</p>',
                                unsafe_allow_html=True
                            )
                            st.title(num_messages)

                        with col2:
                            st.markdown(
                                '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Total Words</p>',
                                unsafe_allow_html=True
                            )
                            st.title(words)

                        with col3:
                            st.markdown(
                                '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Media Shared</p>',
                                unsafe_allow_html=True
                            )
                            st.title(num_media_messages)

                        with col4:
                            st.markdown(
                                '<p style="font-family:Roboto; color:#12ca7d; font-size: 30px; font-weight: bold">Links Shared</p>',
                                unsafe_allow_html=True
                            )
                            st.title(num_links)

                        # Visualization Layout: 2x2 Grid
                        col1, col2 = st.columns([1, 1])  # Equal-width columns

                        with col1:
                            # Monthly Timeline
                            st.markdown(
                                '<u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Monthly Timeline</p></u>',
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This status shows total number of chats per month in group or personal.</p>',
                                unsafe_allow_html=True
                            )
                            timeline = helper.monthly_timeline(selected_user, df)
                            fig = px.line(timeline, x='time', y='message', title='Monthly Chats',
                                          labels={'time': 'Months', 'message': 'Number of Chats'},
                                          template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        with col2:
                            # Daily Timeline
                            st.markdown(
                                '<u><p style="font-family:georgia; color:#ca124d; font-size: 40px; font-weight: bold">Daily Timeline</p></u>',
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This status shows total number of chats per day in group or personal.</p>',
                                unsafe_allow_html=True
                            )
                            daily_timeline = helper.daily_timeline(selected_user, df)
                            fig = px.line(daily_timeline, x='only_date', y='message', title='Daily Chats',
                                          labels={'only_date': 'Date', 'message': 'Number of Chats'},
                                          template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        # Second Row: Most Used Words and Most Active Users
                        col1, col2 = st.columns([1, 1])  # Equal-width columns

                        with col1:
                            # Most Used Words
                            st.markdown(
                                '<u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Used Words</p></u>',
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">Here are the most frequently used words in your chat.</p>',
                                unsafe_allow_html=True
                            )
                            most_used_words_df = get_most_used_words(df)
                            fig = px.bar(most_used_words_df.head(10), x='Word', y='Count',
                                         title='Top 10 Most Used Words',
                                         labels={'Word': 'Words', 'Count': 'Frequency'}, template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        with col2:
                            # Most Active Users
                            if selected_user == 'Overall':
                                st.markdown(
                                    '<u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Active Users</p></u>',
                                    unsafe_allow_html=True
                                )
                                st.markdown(
                                    '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This stats shows the most active users in the group.</p>',
                                    unsafe_allow_html=True
                                )
                                x, new_df = helper.most_active_users(df)
                                fig = px.bar(x, x.index, x.values, title='Top Active User',  # Updated title
                                             labels={'x': 'Users', 'y': 'Number of Chats'}, template='plotly_dark')
                                st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        # Second Part: WordCloud, Emoji Analysis, Sentiment Analysis, Most Active Days

                        # Create a 2x2 grid layout for the visualizations
                        col1, col2 = st.columns([1, 1])  # First row of the grid

                        with col1:
                            # WordCloud
                            st.markdown(
                                '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">WordCloud</p></u>',
                                unsafe_allow_html=True
                            )
                            df_wc = helper.create_wordcloud(selected_user, df)
                            fig = px.imshow(df_wc, title='WordCloud', labels={'x': '', 'y': ''}, template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        with col2:
                            # Emoji Analysis
                            st.markdown(
                                '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Emoji Analysis</p></u>',
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This stats show most used emoji in the chats.</p>',
                                unsafe_allow_html=True
                            )
                            emoji_df = helper.emojis(selected_user, df)
                            col1_emoji, col2_emoji = st.columns([1, 1])  # Sub-grid for emoji analysis

                            with col1_emoji:
                                st.dataframe(emoji_df)

                            with col2_emoji:
                                fig = go.Figure(data=[go.Pie(
                                    labels=emoji_df['Emoji'].head(),  # Adjust column name
                                    values=emoji_df['Count'].head(),  # Adjust column name
                                    hole=0.3
                                )])
                                fig.update_layout(title="Emoji Distribution")
                                st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        # Second Row: Sentiment Analysis and Most Active Days
                        col1, col2 = st.columns([1, 1])  # Second row of the grid

                        with col1:
                            # Sentiment Analysis
                            st.markdown(
                                '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Sentiment Analysis</p></u>',
                                unsafe_allow_html=True
                            )
                            sentiment_counts = analyze_sentiments(df)
                            fig = px.pie(sentiment_counts, values='count', names=sentiment_counts.index,
                                         title='Sentiment Analysis Distribution', template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width

                        with col2:
                            # Most Active Days Bar Chart
                            st.markdown(
                                '<br><br><u><p style="font-family:Roboto; color:#ca124d; font-size: 40px; font-weight: bold">Most Active Days</p></u>',
                                unsafe_allow_html=True
                            )
                            st.markdown(
                                '<p style="font-family:Roboto; color:#ffa81a; font-size: 15px; font-weight: bold">This bar chart displays the days with the highest message count.</p>',
                                unsafe_allow_html=True
                            )
                            most_active_days_df = get_most_active_days(df)
                            fig = px.bar(most_active_days_df.head(10), x='date', y='message_count',
                                         title='Top 10 Most Active Days',
                                         labels={'date': 'Date', 'message_count': 'Message Count'},
                                         template='plotly_dark')
                            st.plotly_chart(fig, use_container_width=True)  # Responsive width
                        # CSV download button
                        csv_report = generate_csv_report(num_messages, words, num_media_messages, num_links,
                                                         most_used_words_df, emoji_df, sentiment_counts)
                        st.download_button(
                            label='Download CSV',
                            data=csv_report,
                            file_name='chat_analysis_report.csv',
                            mime='text/csv'
                        )

                        # Save charts as images
                        save_plot_as_image(px.line(helper.monthly_timeline(selected_user, df), x='time', y='message',
                                                   title='Monthly Chats',
                                                   labels={'time': 'Months', 'message': 'Number of Chats'},
                                                   template='plotly_dark'),
                                           'monthly_timeline.png')

                        save_plot_as_image(px.line(helper.daily_timeline(selected_user, df), x='only_date', y='message',
                                                   title='Daily Chats',
                                                   labels={'only_date': 'Date', 'message': 'Number of Chats'},
                                                   template='plotly_dark'),
                                           'daily_timeline.png')

                        save_plot_as_image(px.bar(get_most_used_words(df).head(10), x='Word', y='Count',
                                                  title='Top 10 Most Used Words',
                                                  labels={'Word': 'Words', 'Count': 'Frequency'},
                                                  template='plotly_dark'),
                                           'most_used_words.png')

                        save_plot_as_image(
                            go.Figure(data=[go.Pie(labels=helper.emojis(selected_user, df)['Emoji'].head(),
                                                   values=helper.emojis(selected_user, df)['Count'].head(), hole=0.3)]),
                            'emoji_analysis.png')

                        save_plot_as_image(
                            px.pie(analyze_sentiments(df), values='count', names=analyze_sentiments(df).index,
                                   title='Sentiment Analysis Distribution', template='plotly_dark'),
                            'sentiment_analysis.png')

                        save_plot_as_image(px.bar(get_most_active_days(df).head(10), x='date', y='message_count',
                                                  title='Top 10 Most Active Days',
                                                  labels={'date': 'Date', 'message_count': 'Message Count'},
                                                  template='plotly_dark'),
                                           'most_active_days.png')

                        # Generate Word report
                        word_report = generate_word_report(num_messages, words, num_media_messages, num_links,
                                                           most_used_words_df, emoji_df, sentiment_counts,
                                                           'monthly_timeline.png', 'daily_timeline.png',
                                                           'most_used_words.png',
                                                           'emoji_analysis.png', 'sentiment_analysis.png',
                                                           'most_active_days.png')

                        # Provide the download button for the Word report
                        st.download_button(
                            label='Download Word',
                            data=word_report,
                            file_name='chat_analysis_report.docx',
                            mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        )


if __name__ == "__main__":
    st.set_page_config(layout="wide")

    main()
